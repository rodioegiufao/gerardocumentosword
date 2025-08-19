import streamlit as st
import pandas as pd
from docx import Document
import os
from datetime import datetime
from io import BytesIO
import base64

# Configuração da página
st.set_page_config(
    page_title="Gerador de Documentos Técnicos",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Mapeamento dos meses em português
MESES_PT_BR = {
    "January": "janeiro",
    "February": "fevereiro",
    "March": "março",
    "April": "abril",
    "May": "maio",
    "June": "junho",
    "July": "julho",
    "August": "agosto",
    "September": "setembro",
    "October": "outubro",
    "November": "novembro",
    "December": "dezembro"
}

# Dados fixos
ENGENHEIROS = {
    "SALOMÃO JOSE COHEN": {
        "CREA": "0401863549",
        "EMAIL": "salomao.cohen@hotmail.com",
        "FONE": "(92) 99136-1006",
        "ENDERECO": "Rua Mar de SUFE, 67, Conjunto Imperial, Flores, Manaus/AM, CEP 69058-438",
        "RG": "801.420-5",
        "CPF": "317.323.132-53"
    },
    "RODRIGO DAMASCENO NASCIMENTO": {
        "CREA": "0920192912",
        "EMAIL": "rodrigo.ele@ribeirolopes.eng.br",
        "FONE": "(95) 99146-6367",
        "ENDERECO": "Rua Antonio Marques, 108, Buritis, Boa Vista/RR, CEP 69309-172",
        "RG": "413.816-3",
        "CPF": "022.331.622-93"
    }
}

CLIENTES = {
    "SUPERINTENDÊNCIA DA RECEITA FEDERAL": {
        "CNPJ": "00.394.460/0070-73",
        "END_SEDE": "Rua Travessa Rui Barbosa,1039, Bairro Reduto, Belém-PA, CEP 66.053-260",
        "RESPONSAVEL": "Eduardo Badaró Fernandes",
        "NACIONALIDADE": "brasileira",
        "ESTADO_CIVIL": "Solteiro(a)",
        "IDENTIDADE": "01.184.275-1",
        "CPF": "000.934.647-38",
        "DOMICILIO": "Av. Dr. Theomario Pinto da Costa, n° 192, Chapada, Manaus/AM"
    },
    "INSTITUTO DE PREVIDÊNCIA SOCIAL DO ESTADO DE RORAIMA - IPER": {
        "CNPJ": "03.491.063/0001-86",
        "END_SEDE": "Rua Araújo Filho, 823 - Centro, Boa Vista - RR, CEP: 69.301-090",
        "RESPONSAVEL": "Rafael David Aires Alencar",
        "NACIONALIDADE": "brasileiro",
        "ESTADO_CIVIL": "Casadp(a)",
        "IDENTIDADE": "15.912-4",
        "CPF": " 512.997.122-15",
        "DOMICILIO": "Rua Araújo Filho, 823 - Centro, Boa Vista - RR, CEP: 69.301-090"
    },
    "SECRETARIA MUNICIPAL DE OBRAS - SMO": {
        "CNPJ": "05.943.030/0001-55",
        "END_SEDE": "Av. Santos Dumont, 1721 - São Francisco, Boa Vista - RR, CEP 69.305-105",
        "RESPONSAVEL": "Kaynara Carvalho de Oliveira",
        "NACIONALIDADE": "brasileira",
        "ESTADO_CIVIL": "Solteiro(a)",
        "IDENTIDADE": "361.182-7",
        "CPF": "062.137.393-19",
        "DOMICILIO": "Rua João XXIII, n° 476, apartamento 4, Bairro Aparecida, Boa Vista/RR"
    },
    "SECRETARIA DE ESTADO DE INFRAESTRUTURA DE RORAIMA - SEINF": {
        "CNPJ": "84.012.012/0001-26",
        "END_SEDE": "Av. Getúlio Vargas,3941 Bairro Canarinho, Boa Vista/Roraima, CEP 69.306-545",
        "RESPONSAVEL": "Raissa Karla Santos de Andrade",
        "NACIONALIDADE": "brasileira",
        "ESTADO_CIVIL": "Solteiro(a)",
        "IDENTIDADE": "202.909-6",
        "CPF": "049.666.684-33",
        "DOMICILIO": "Av. Getúlio Vargas,3941 Bairro Canarinho, Boa Vista/Roraima, CEP 69.306-545"
    },
    "SECRETARIA DE ESTADO DA SAÚDE DE RORAIMA - SESAU": {
        "CNPJ": "84.013.408/0001-98",
        "END_SEDE": "Rua Madri, 180, Aeroporto, Boa Vista - RR, CEP 69.310-043",
        "RESPONSAVEL": "Wengley Glides Martins Silva",
        "NACIONALIDADE": "brasileiro",
        "ESTADO_CIVIL": "Solteiro(o)",
        "IDENTIDADE": "15.212-3",
        "CPF": "779.133.102-00",
        "DOMICILIO": "Rua Olavo Brasil filho, 101, apartamento 03, Jardim Floresta, CEP 69.312-133"
    }
}

def extrair_municipio(endereco_completo):
    """Extrai o município de um endereço completo."""
    partes = endereco_completo.split(',')
    if len(partes) >= 3:
        cidade_uf = partes[-1].strip()
        cidade_uf = cidade_uf.split('CEP')[0].strip()
        if '-' in cidade_uf:
            cidade = cidade_uf.split('-')[0].strip()
            return cidade
    return "Boa Vista"

def verificar_templates():
    """Verifica se os templates necessários existem na pasta atual."""
    templates = {
        "Memorial descritivo - Teste.docx": os.path.join(os.path.dirname(__file__), "Memorial descritivo - Teste.docx"),
        "PROCURAÇÃO - TESTE.docx": os.path.join(os.path.dirname(__file__), "PROCURAÇÃO - TESTE.docx"),
        "fazer-termo-de-responsabilidade-para-uso-geracao-propria-pessoa-fisica -TESTE.docx": os.path.join(os.path.dirname(__file__), "fazer-termo-de-responsabilidade-para-uso-geracao-propria-pessoa-fisica -TESTE.docx"),
        "CARTA VIABILIDADA - TESTE.docx": os.path.join(os.path.dirname(__file__), "CARTA VIABILIDADA - TESTE.docx"),
        "Termo de responsabilidade de não Utilização de Geração Própria - TESTE.docx": os.path.join(os.path.dirname(__file__), "Termo de responsabilidade de não Utilização de Geração Própria - TESTE.docx")
    }
    
    templates_encontrados = {}
    for nome, caminho in templates.items():
        if os.path.exists(caminho):
            templates_encontrados[nome] = caminho
        else:
            st.warning(f"Template não encontrado: {caminho}")
    
    return templates_encontrados

def processar_documento(arquivo_origem, dados):
    """Processa um documento Word substituindo os placeholders."""
    try:
        doc = Document(arquivo_origem)
        
        # Substituir texto em parágrafos
        for paragraph in doc.paragraphs:
            for chave, valor in dados.items():
                if chave in paragraph.text:
                    paragraph.text = paragraph.text.replace(chave, str(valor))
        
        # Substituir texto em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for chave, valor in dados.items():
                        if chave in cell.text:
                            cell.text = cell.text.replace(chave, str(valor))
        
        # Salvar em buffer de memória
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    except Exception as e:
        st.error(f"Erro ao processar {arquivo_origem}: {str(e)}")
        return None

def criar_link_download(buffer, nome_arquivo, texto_link):
    """Cria um link para download do arquivo."""
    if buffer:
        b64 = base64.b64encode(buffer.getvalue()).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="{nome_arquivo}">{texto_link}</a>'
        return href
    return None

def main():
    # Cabeçalho com estilo melhorado
    st.markdown("""
        <style>
        .main-header {
            font-size: 3rem;
            color: #1f77b4;
            text-align: center;
            margin-bottom: 1rem;
        }
        .sub-header {
            font-size: 1.5rem;
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 0.5rem;
            margin-top: 1.5rem;
        }
        .info-box {
            background-color: #f8f9fa;
            border-left: 4px solid #3498db;
            padding: 1rem;
            border-radius: 0.25rem;
            margin-bottom: 1rem;
        }
        .download-section {
            background-color: #e8f4f8;
            padding: 1.5rem;
            border-radius: 0.5rem;
            margin-top: 1.5rem;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<h1 class="main-header">📄 Gerador de Documentos Técnicos</h1>', unsafe_allow_html=True)
    st.markdown("---")
    
    # Verificar templates disponíveis
    templates = verificar_templates()
    
    if not templates:
        st.error("Nenhum template encontrado. Certifique-se de que os arquivos estão na mesma pasta do aplicativo.")
        st.info("""
        Templates necessários:
        - Memorial descritivo - Teste.docx
        - PROCURAÇÃO - TESTE.docx
        - fazer-termo-de-responsabilidade-para-uso-geracao-propria-pessoa-fisica -TESTE.docx
        - CARTA VIABILIDADA - TESTE.docx
        - Termo de responsabilidade de não Utilização de Geração Própria - TESTE.docx
        """)
        return
    
    # Sidebar com informações
    with st.sidebar:
        st.markdown("### 📋 Templates Disponíveis")
        for nome, caminho in templates.items():
            st.success(f"✅ {nome}")
        
        st.markdown("---")
        st.markdown("### ℹ️ Instruções")
        st.info("""
        1. Preencha todos os campos do formulário
        2. Selecione o engenheiro e cliente
        3. Clique em 'Gerar Documentos'
        4. Faça o download dos documentos gerados
        """)
        
        st.markdown("---")
        st.markdown("### 🔧 Suporte")
        st.info("Em caso de problemas, entre em contato com o administrador do sistema.")
    
    # Formulário de dados do projeto
    with st.form("dados_projeto"):
        st.markdown('<h2 class="sub-header">📝 Dados do Projeto</h2>', unsafe_allow_html=True)
        
        # Informações básicas
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 🔌 Dados Elétricos")
            potencia = st.text_input("POTÊNCIA DA SUBESTAÇÃO EM KVA*", placeholder="Ex: 225")
            art = st.text_input("ART*", placeholder="Ex: AM20240488640")
            tensao = st.text_input("TENSÃO GERAL*", placeholder="Ex: 380/220, 220/127V")
            ramal_tamanho = st.text_input("Tamanho do ramal de entrada*", placeholder="Ex: 150")
            ramal_cabo = st.text_input("Cabo do ramal de entrada MT*", placeholder="Ex: 35mm²")
            
        with col2:
            st.markdown("#### 📋 Informações Gerais")
            nome_projeto = st.text_input("Nome do projeto*", placeholder="Ex: Construção de Escola")
            concessionaria = st.text_input("Concessionária*", placeholder="Ex: RORAIMA ENERGIA", value="RORAIMA ENERGIA")
            numero_uc = st.text_input("Número da UC*", placeholder="Ex: UCXXXXXXX ou S/N°")
            demanda = st.text_input("Demanda do Projeto*", placeholder="Ex: 1500")
        
        # Localização
        st.markdown("#### 📍 Localização")
        endereco_empreendimento = st.text_input("ENDEREÇO DO EMPREENDIMENTO*", 
                                               placeholder="Ex: Rua Coronel Pinto, nº 664, Centro, Boa Vista - RR")
        localizacao_projeto = st.text_input("Coordenadas Geográficas*", 
                                           placeholder="Ex: 2°50'29.1'N 60°42'27.4'W")
        
        # Datas
        st.markdown("#### 📅 Datas da Obra")
        col3, col4 = st.columns(2)
        with col3:
            data_inicio = st.text_input("Início da Obra*", placeholder="Ex: 10/08/2025")
        with col4:
            data_fim = st.text_input("Fim da Obra*", placeholder="Ex: 10/08/2065")
        
        # Seleção de engenheiro e cliente
        st.markdown('<h2 class="sub-header">👥 Seleção de Profissionais e Clientes</h2>', unsafe_allow_html=True)
        
        col5, col6 = st.columns(2)
        
        with col5:
            engenheiro = st.selectbox("Engenheiro Responsável*", list(ENGENHEIROS.keys()))
        
        with col6:
            cliente = st.selectbox("Cliente*", list(CLIENTES.keys()))
        
        # Opção para gerar termo de não geração
        st.markdown('<h2 class="sub-header">⚡ Opções de Documentos</h2>', unsafe_allow_html=True)
        gerar_termo_nao_geracao = st.checkbox("Gerar Termo de Não Utilização de Geração Própria", value=True)
        
        # Botão de envio
        submitted = st.form_submit_button("🔄 Gerar Documentos", use_container_width=True)
    
    if submitted:
        # Validar campos obrigatórios
        campos_obrigatorios = [
            potencia, art, tensao, ramal_tamanho, ramal_cabo, 
            nome_projeto, concessionaria, endereco_empreendimento, 
            localizacao_projeto, numero_uc, demanda, data_inicio, data_fim
        ]
        
        if not all(campos_obrigatorios):
            st.error("Por favor, preencha todos os campos obrigatórios (*)")
            return
        
        # Preparar dados
        dados = {
            'XXXX': potencia,
            'YYYY': art,
            'XXXY': endereco_empreendimento,
            'DDDD': tensao,
            'EEEE': ramal_tamanho,
            'FFFF': ramal_cabo,
            'GGGG': nome_projeto,
            'HHHH': localizacao_projeto,
            'LLLL': concessionaria,
            'VVVV': numero_uc,
            'ZXZX': demanda,
            'DTIN': data_inicio,
            'DTFI': data_fim,
            'ZZZZ': extrair_municipio(localizacao_projeto)
        }
        
        # Adicionar dados do engenheiro
        dados['XXYY'] = engenheiro
        eng_data = ENGENHEIROS[engenheiro]
        dados['AAAA'] = eng_data["CREA"]
        dados['BBBB'] = eng_data["EMAIL"]
        dados['CCCC'] = eng_data["FONE"]
        dados['IIII'] = eng_data["RG"]
        dados['JJJJ'] = eng_data["CPF"]
        dados['KKKK'] = eng_data["ENDERECO"]
        
        # Adicionar dados do cliente
        dados['MMMM'] = cliente
        cli_data = CLIENTES[cliente]
        dados['NNNN'] = cli_data["CNPJ"]
        dados['OOOO'] = cli_data["END_SEDE"]
        dados['PPPP'] = cli_data["RESPONSAVEL"]
        dados['QQQQ'] = cli_data["NACIONALIDADE"]
        dados['RRRR'] = cli_data["ESTADO_CIVIL"]
        dados['SSSS'] = cli_data["IDENTIDADE"]
        dados['TTTT'] = cli_data["CPF"]
        dados['UUUU'] = cli_data["DOMICILIO"]
        
        # Adicionar datas formatadas
        hoje = datetime.now()
        dados['14 / 07 / 2026'] = hoje.strftime('%d / %m / %Y')
        
        mes_ingles = hoje.strftime('%B')
        mes_portugues = MESES_PT_BR.get(mes_ingles, mes_ingles)
        dados['14 de julho de 2025'] = hoje.strftime(f'%d de {mes_portugues} de %Y')
        dados['MAIO / 2025'] = hoje.strftime(f'{mes_portugues.upper()} / %Y')
        
        # Processar documentos
        documentos_gerados = {}
        
        # Nomes dos templates
        template_names = {
            "Memorial Descritivo": "Memorial descritivo - Teste.docx",
            "Procuração": "PROCURAÇÃO - TESTE.docx",
            "Termo de Responsabilidade": "fazer-termo-de-responsabilidade-para-uso-geracao-propria-pessoa-fisica -TESTE.docx",
            "Carta de Viabilidade": "CARTA VIABILIDADA - TESTE.docx"
        }
        
        # Adicionar termo de não geração se selecionado
        if gerar_termo_nao_geracao:
            template_names["Termo de Não Geração"] = "Termo de responsabilidade de não Utilização de Geração Própria - TESTE.docx"
        
        for doc_name, template_file in template_names.items():
            if template_file in templates:
                buffer = processar_documento(templates[template_file], dados)
                if buffer:
                    if doc_name == "Termo de Responsabilidade":
                        nome_arquivo = f"Termo de Responsabilidade - {cli_data['RESPONSAVEL']}.docx"
                    elif doc_name == "Carta de Viabilidade":
                        nome_arquivo = f"Carta de Viabilidade - {nome_projeto}.docx"
                    elif doc_name == "Termo de Não Geração":
                        nome_arquivo = f"Termo de Não Utilização de Geração Própria - {nome_projeto}.docx"
                    else:
                        nome_arquivo = f"{doc_name} - {nome_projeto}.docx"
                    
                    documentos_gerados[doc_name] = {
                        "buffer": buffer,
                        "nome": nome_arquivo
                    }
        
        # Exibir resultados
        if documentos_gerados:
            st.success("✅ Documentos gerados com sucesso!")
            st.markdown('<h2 class="sub-header">📥 Download dos Documentos</h2>', unsafe_allow_html=True)
            
            # Seção de download com layout melhorado
            st.markdown('<div class="download-section">', unsafe_allow_html=True)
            
            # Organizar os documentos em colunas
            cols = st.columns(min(3, len(documentos_gerados)))
            
            for i, (doc_name, doc_info) in enumerate(documentos_gerados.items()):
                with cols[i % len(cols)]:
                    st.markdown(f"##### {doc_name}")
                    st.markdown(f"*Arquivo: {doc_info['nome']}*")
                    
                    # Botão de download
                    st.download_button(
                        label=f"⬇️ Baixar {doc_name}",
                        data=doc_info["buffer"].getvalue(),
                        file_name=doc_info["nome"],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_{i}"
                    )
            
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Opção para baixar todos os documentos em um ZIP
            st.markdown("---")
            st.markdown("#### 📦 Baixar Todos os Documentos")
            
            # Criar arquivo ZIP com todos os documentos
            import zipfile
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for doc_name, doc_info in documentos_gerados.items():
                    zip_file.writestr(doc_info['nome'], doc_info['buffer'].getvalue())
            
            zip_buffer.seek(0)
            
            st.download_button(
                label="📦 Baixar Todos os Documentos (ZIP)",
                data=zip_buffer.getvalue(),
                file_name=f"Documentos_{nome_projeto}.zip",
                mime="application/zip",
                key="download_all"
            )
            
            # Mostrar resumo dos dados
            with st.expander("📋 Visualizar Dados Preenchidos"):
                dados_para_exibir = {k: v for k, v in dados.items() if not k.startswith('14') and k != 'MAIO / 2025'}
                st.json(dados_para_exibir)
        else:
            st.error("❌ Nenhum documento foi processado. Verifique se os templates estão corretos.")

if __name__ == "__main__":

    main()

